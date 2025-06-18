
import os
import sys
import subprocess

# Auto-install required packages
def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

# Importing necessary packages
try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    install("python-docx")
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

try:
    import pdfplumber
except ImportError:
    install("pdfplumber")
    import pdfplumber

# Function to add right-aligned page numbers in the footer
def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = 2  # 0=Left, 1=Center, 2=Right (Right-aligned), depends on how you want to insert page number
    run = paragraph.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = " PAGE "

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

# Function to extract text from PDF using pdfplumber
def extract_text_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

# Main function to split the transcripts
def split_transcripts(input_path, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    ext = os.path.splitext(input_path)[1].lower()

    if ext == ".docx":
        doc = Document(input_path)
        paragraphs = [para.text for para in doc.paragraphs]
    elif ext == ".pdf":
        raw_text = extract_text_from_pdf(input_path)
        paragraphs = raw_text.split("\n")
    else:
        raise ValueError("Unsupported file type. Please use a .docx or .pdf file.")

    current_doc = None
    current_name = None

    for text in paragraphs:
        if text.strip().startswith("## Respondent"):  ## ("## Respondent")This is the indicator of what's the heading should be. Feel free to change accordingly.
            if current_doc and current_name:
                filename = os.path.join(output_dir, f"{current_name}.docx")
                current_doc.save(filename)

            current_doc = Document()
            section = current_doc.sections[0]
            add_page_number(section)

            current_name = text.strip().replace("## ", "").replace(":", "").replace(" ", "_")
            current_doc.add_paragraph(text)
        else:
            if current_doc:
                current_doc.add_paragraph(text)

    if current_doc and current_name:
        filename = os.path.join(output_dir, f"{current_name}.docx")
        current_doc.save(filename)

# Example usage
# Input_file is the name of the docuemnt that you want to split
# Output folder is the name of folder where you wnat to store splited document. Please change it accordingly.
if __name__ == "__main__":
    input_file = "transcripts.docx" 
    output_folder = "split_transcripts"
    split_transcripts(input_file, output_folder)
